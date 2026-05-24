import streamlit as st
import os
import io
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image

# Safely import python-docx for Word generation/reading
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Safely import PyPDF2 for reading PDF CVs
try:
    import PyPDF2
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# 1. PAGE SETUP
st.set_page_config(page_title="EliteConnect AI", layout="centered")
load_dotenv()

# API Configuration
api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)

# Initialize Session State for Dynamic Skills
if "user_skills" not in st.session_state:
    st.session_state.user_skills = []

# --- CALLBACK FUNCTION FOR ENTER KEY ---
def add_skill():
    skill = st.session_state.skill_input.strip()
    if skill and skill not in st.session_state.user_skills:
        st.session_state.user_skills.append(skill)
    # Clear input after pressing enter or add button
    st.session_state.skill_input = "" 
    # Expertise Section display fix
if st.session_state.user_skills:
    st.info("Added Skills: " + ", ".join(st.session_state.user_skills))

# 2. CUSTOM CSS (Premium White Background Inputs + Perfect Alignment)
st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, rgba(255, 20, 147, 0.8), rgba(255, 255, 255, 0.95), rgba(0, 191, 255, 0.8));
        background-attachment: fixed;
    }
    .magic-title { 
        text-align: center; 
        font-size: 3.5rem; 
        font-weight: 900; 
        background: linear-gradient(to right, #d1005a, #0044b3); 
        -webkit-background-clip: text; 
        -webkit-text-fill-color: transparent; 
        text-shadow: 3px 3px 8px rgba(0,0,0,0.2); 
        margin-bottom: 0px; 
        font-family: 'Arial', sans-serif; 
    }
    .magic-subtitle { 
        text-align: center; 
        font-size: 1.2rem; 
        color: #222222; 
        font-weight: 600; 
        margin-bottom: 40px; 
    }

    /* PREMIUM SOLID WHITE BACKGROUND FOR ALL INPUTS */
    div[data-baseweb="input"] > div, 
    div[data-baseweb="textarea"] > div, 
    div[data-baseweb="select"] > div,
    div[data-testid="stFileUploadDropzone"] {
        background-color: #ffffff !important; 
        border: 1px solid #d1d5db !important;
        border-radius: 10px !important;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05) !important;
        transition: all 0.3s ease !important;
    }
    
    /* Input Focus Hover Effects */
    div[data-baseweb="input"] > div:focus-within, 
    div[data-baseweb="textarea"] > div:focus-within,
    div[data-baseweb="select"] > div:focus-within {
        border: 1px solid #0044b3 !important;
        box-shadow: 0 4px 12px rgba(0, 68, 179, 0.15) !important;
    }

    label { 
        font-weight: 700 !important; 
        color: #111111 !important; 
        margin-bottom: 5px !important;
    }
    
    /* Button Styling to match premium look */
    div[data-testid="stButton"] button {
        border-radius: 8px !important;
        font-weight: bold !important;
    }
    </style>
""", unsafe_allow_html=True)

# --- HEADER ---
st.markdown('<h1 class="magic-title">EliteConnect AI</h1>', unsafe_allow_html=True)
st.markdown('<p class="magic-subtitle">Architect top-tier proposals and emails that command authority and close deals.</p>', unsafe_allow_html=True)

# ==========================================
# STEP 1: EXPERTISE (DYNAMIC TYPING)
# ==========================================
st.markdown("### Step 1: Tell Us About Your Expertise")

col_input, col_btn = st.columns([3, 1], vertical_alignment="bottom")
with col_input:
    st.text_input("Type your skill and add:", placeholder="e.g., React Native, Video Editing...", autocomplete="off", key="skill_input", on_change=add_skill)
with col_btn:
    st.button("Add ➕", use_container_width=True, on_click=add_skill)

if st.session_state.user_skills:
    st.write("Your Skills: " + ", ".join([f"**{s}**" for s in st.session_state.user_skills]))

freelancer_text = st.text_area("Detail your portfolio highlights or Instructions:", placeholder="e.g., Keep the response brief. Focus on my recent dashboard project...", height=100, key="freelancer_text")
freelancer_files = st.file_uploader("Upload portfolio/CV files (PDF, DOCX, TXT, Images):", accept_multiple_files=True, key="freelancer_files_uploader")

st.markdown("---")

# ==========================================
# STEP 2: CLIENT REQUIREMENTS
# ==========================================
st.markdown("### Step 2: Define the Client's Needs")
client_category = st.selectbox("Project category:", ["Social Media Video/Content", "App/Website Design", "Branding & Logo", "Software Development", "Writing & Translation", "Other"], key="client_cat")
client_text = st.text_area("Paste the job description:", placeholder="e.g., Need a YouTube editor for a soccer channel.", height=120, key="client_text_area")
client_files = st.file_uploader("Upload client reference files (Images/Screenshots):", accept_multiple_files=True, type=['png', 'jpg', 'jpeg', 'webp'], key="client_files_uploader")

st.markdown("---")

# ==========================================
# STEP 3: OUTPUT & STRATEGY
# ==========================================
st.markdown("### Step 3: Output & Strategy Configuration")
output_type = st.selectbox("What do you want to generate?", ["Project Proposal", "Client Email"], key="out_type")
tone = st.selectbox("Select the strategic tone:", ["Student / Entry-Level (Humble, eager, hardworking)", "To the point and Professional (Executive level)", "Friendly and Conversational (Relationship-building)", "Highly Confident (Industry authority)"], key="tone_type")
strategy_uvp = st.selectbox("Strategic Advantage:", ["AUTO-DETECT CLIENT PSYCHOLOGY", "The Competitor Takedown", "The Risk Reversal", "The Visionary Pitch"], key="uvp_type")

available_models = ["models/gemini-1.5-flash"] 
if api_key:
    try:
        fetched_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        if fetched_models: available_models = fetched_models
    except Exception: pass
selected_model = st.selectbox("Select AI Engine:", available_models, key="ai_model")

# Generate Logic
if st.button("Generate Content", type="primary", use_container_width=True):
    if not api_key: 
        st.error("API Key missing! Please check your .env file.")
    elif not client_text and not client_files: 
        st.warning("⚠️ Please provide a job description text OR upload a reference image/screenshot.")
    else:
        with st.spinner("Analyzing requirements & engineering your customized copy..."):
            try:
                model = genai.GenerativeModel(selected_model)
                skills_str = ", ".join(st.session_state.user_skills) if st.session_state.user_skills else "General technical skills"
                
                prompt_text = client_text if client_text else "the requirements mentioned in the attached images"
                
                # Extract text from CV/portfolio files safely
                freelancer_extracted_text = ""
                media_files = [] 
                
                if freelancer_files:
                    for f in freelancer_files:
                        ext = f.name.split('.')[-1].lower()
                        if ext == 'pdf' and HAS_PYPDF:
                            try:
                                pdf_reader = PyPDF2.PdfReader(f)
                                for page in pdf_reader.pages:
                                    text = page.extract_text()
                                    if text: freelancer_extracted_text += text + "\n"
                            except: pass
                        elif ext in ['docx', 'doc'] and HAS_DOCX:
                            try:
                                doc = Document(f)
                                for para in doc.paragraphs:
                                    freelancer_extracted_text += para.text + "\n"
                            except: pass
                        elif ext == 'txt':
                            try:
                                freelancer_extracted_text += f.getvalue().decode("utf-8") + "\n"
                            except: pass
                        elif ext in ['png', 'jpg', 'jpeg', 'webp']:
                            try:
                                media_files.append(Image.open(f))
                            except: pass
                
                if client_files:
                    for cf in client_files:
                        try:
                            media_files.append(Image.open(cf))
                        except: pass

                # ---------------------------------------------------------
                # Fully Redesigned Global Adaptive Prompt
                # ---------------------------------------------------------
                full_prompt = f"""
                You are a world-class sales copywriter and enterprise growth strategist. Your task is to generate a premium, production-ready document tailored explicitly to the requested parameters. 

                --- [URGENT: UNIVERSAL LANGUAGE & HOLISTIC ASSESSMENT DIRECTIVE] ---
                1. UNIVERSAL LANGUAGE SUPPORT: Analyze the inputs below for constraints, instructions, or descriptions written in ANY language (e.g., English, Roman Urdu, Spanish, French, Hindi, etc.). You MUST generate the final output in the language explicitly requested by the user. If no specific language is requested, seamlessly match the primary language used in the 'Job Description' or 'User Manual Input'.
                2. HOLISTIC INPUT ASSESSMENT: You must synthesize and respect EVERY single input provided. Every dropdown selection, every typed instruction, every extracted CV detail, every selected skill, and the content of every uploaded image must influence the final result. Do not ignore any user constraint or provided data point.
                3. STRICT LENGTH CONSTRAINTS: If the user specifies a length limit (e.g., 'short', 'brief', 'chota', '5 lines', 'solo 3 lineas'), you must compress or shorten ONLY the body paragraphs/explanations. The overall formal structure and structural checkpoints must NEVER be skipped or compromised.

                --- CRITICAL FORMATTING GUARANTEE (NON-NEGOTIABLE) ---
                You MUST strictly follow the structural schema below based on the chosen output type:

                IF OUTPUT TYPE IS 'Client Email':
                - Subject Line: [Craft a highly contextual, high-open-rate subject line reflecting the Strategy/UVP chosen]
                - Salutation: [Formal greetings matching the selected Tone, e.g., 'Hi [Name],' or 'Dear [Team],']
                - Hook: [An immediate, gripping problem-centric opening paragraph]
                - Core Pitch / Body Paragraphs: [The core value mapping. Infuse the Strategic Tone and Strategic Advantage deeply here]
                - Call to Action (CTA): [A single, friction-free clear next step request]
                - Sign-off: [Professional validation and a clean, placeholder-free signature block]

                IF OUTPUT TYPE IS 'Project Proposal':
                - Header Title: [Professional Document Title, e.g., 'Technical & Creative Project Proposal: [Project Title]']
                - Executive Summary: [A high-level diagnostic summary showcasing deep understanding of the client's current pain points]
                - Scope of Work & Implementation Solution: [Granular, well-structured structural layout or bulleted list of what will be built/delivered]
                - Strategic Advantage / UVP: [Deep explanation of why this solution wins, pivoting hard on the selected Strategic Advantage]
                - Deliverables & High-Level Timeline: [Clear checkpoints mapping the output category requirements]
                - Next Steps & Call to Action: [Clean onboarding pathway instructions]
                - Professional Sign-off: [Signature block without bracketed variables]

                --- TRUE STRATEGIC DIRECTIONAL PIVOTING RULES ---
                Change the fundamental angle, vocabulary, and persuasive psychology of the copy based on these selections:

                1. Selected STRATEGIC TONE: '{tone}'
                   - 'Student / Entry-Level': Focus heavily on unparalleled work ethic, high coachability, dedicated research, and extreme diligence.
                   - 'To the point and Professional': Highly analytical, crisp sentences, metrics-oriented, zeros out fluff, focuses strictly on ROI and executive value.
                   - 'Friendly and Conversational': Warm, relational, focuses on seamless collaborative chemistry, open feedback loops, and building long-term partnership.
                   - 'Highly Confident': Bold, authoritative, uses premium industry vernacular, positions you as the definitive expert who dictates the path forward.

                2. Selected STRATEGIC ADVANTAGE / UVP Angle: '{strategy_uvp}'
                   - 'AUTO-DETECT CLIENT PSYCHOLOGY': Diagnose hidden risks, bottlenecks, or implicit fears suggested in the job description or context files, then solve them.
                   - 'The Competitor Takedown': Focus on how standard freelancers take shortcuts or overcharge, showing why your specific execution blueprint outperforms the status quo.
                   - 'The Risk Reversal': De-risk the arrangement entirely (e.g., offering transparent milestones, preliminary mockups, or clear delivery guarantees).
                   - 'The Visionary Pitch': Paint a vivid picture of macro-scale success, focus on exponential value, scalability, and long-term positioning.

                --- CONTEXT DATA & USER SPECIFICATIONS ---
                * Target Structure Type: {output_type}
                * Project Category Context: {client_category}
                * Job Description / Client Pain Points: {prompt_text}
                * Highlighted Core Competencies / Skills: {skills_str}
                * User Custom Instructions (Prioritize length, language, and specific demands here): {freelancer_text if freelancer_text else 'N/A'}
                * Extracted Background Profiling Text (from CV/Files): {freelancer_extracted_text if freelancer_extracted_text else 'None'}

                --- OUTPUT SANITIZATION CHECK ---
                - Do NOT leave empty square brackets or placeholders like '[Your Name]', '[Client Name]', or '[Insert Date]'. If a variable is missing from the context background, write a completely natural, flowing sentence that requires zero placeholders.
                - Output ONLY the clean, ready-to-use document content. Do not provide meta-commentary or conversational intros like 'Here is your proposal:'.
                """
                
                content_to_send = [full_prompt]
                content_to_send.extend(media_files) 
                
                response = model.generate_content(content_to_send)
                
                st.session_state.final_magic_text = response.text
                st.session_state.current_output_type = output_type
            
            # --- INTELLIGENT DUAL-QUOTA ERROR HANDLING BLOCK ---
            except Exception as e: 
                error_message = str(e)
                if "429" in error_message or "Quota" in error_message or "quota" in error_message.lower():
                    import re
                    
                    # 1. Check if the error is explicitly a Daily Quota Limit exhaustion
                    if "GenerateRequestsPerDay" in error_message or "daily" in error_message.lower():
                        st.error("⏳ Daily API Quota Exhausted! You have reached the maximum allowed free requests for today. Please try again tomorrow or change your API key.")
                    
                    # 2. Check if there are short-term retry seconds
                    else:
                        match = re.search(r"retry in\s+([0-9.]+)\s*s", error_message, re.IGNORECASE)
                        if match:
                            wait_seconds = int(float(match.group(1))) + 1
                            # If Google sends a giant retry number, it's also effectively a daily/long block
                            if wait_seconds > 120:
                                st.error(f"⏳ Daily API Quota Exhausted! Please wait for the daily reset or change your API key (Estimated wait: {wait_seconds // 60} minutes).")
                            else:
                                st.error(f"⏳ API Rate Limit Reached! Please wait exactly {wait_seconds} seconds before clicking 'Generate' again.")
                        else:
                            st.error("⏳ API Rate Limit Reached! Free tier limit exceeded. Please wait a few moments and try again.")
                else:
                    st.error(f"⚠️ Error during generation: {error_message}")

# Display Results
if "final_magic_text" in st.session_state:
    st.markdown("---")
    st.markdown("### ✨ Your Generated Content")
    st.text_area("Review and Copy:", value=st.session_state.final_magic_text, height=350)
    
    if HAS_DOCX:
        doc = Document()
        doc.add_paragraph(st.session_state.final_magic_text)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        dynamic_file_name = f"{st.session_state.current_output_type.replace(' ', '_')}.docx"
        
        st.download_button(
            label=f"DOWNLOAD AS WORD DOCUMENT (.docx)", 
            data=buffer, 
            file_name=dynamic_file_name, 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            use_container_width=True
        )
    else:
        st.error("⚠️ Library 'python-docx' is not installed. Please run `pip install python-docx` in your terminal to enable downloading.")